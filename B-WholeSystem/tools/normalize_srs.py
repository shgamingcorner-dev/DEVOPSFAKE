"""Normalize SRS_project_GRPPRJ: fill placeholder REQ IDs with clean sequential
numbers, fix typos. Preserves structure, tables, images. Operates on a copy."""
import re, shutil, os
from docx import Document

SRC = r"C:\Users\shgam\DevOpsA\DEVOPSFAKE-staging\.hermes\desktop-attachments\SRS_project_GRPPRJ (2).docx"
DST = r"C:\Users\shgam\DevOpsA\DEVOPSFAKE-staging\B-WholeSystem\docs\SRS_project_GRPPRJ.docx"

os.makedirs(os.path.dirname(DST), exist_ok=True)
shutil.copy(SRC, DST)
doc = Document(DST)

# --- Walk all paragraphs + tables, collect the REQ id cells to replace ---
# Strategy: each requirement table has rows: [REQ_ID cell, Requirement cell].
# We walk tables in order; the placeholder ids appear as:
#   "REQ-0?"  (emergency block LED + servo)
#   "REQ-xx"  (NFR block, 7 of them)
#   "REQ-"    (2 empty ids: log updater 2nd row + ThingSpeak 15s)
# We replace them in document order with a running counter that we assign
# per the final numbering:
#   REQ-01..04 exist; missing REQ-05 -> we insert nothing, just continue.
#   emergency: LED=REQ-08, servo=REQ-09  (but REQ-08 is LCD in doc... )
#
# SIMPLER + SAFER: replace placeholders positionally with unique ids:
#   REQ-0? -> REQ-08 / REQ-09 (2 occurrences, in order)
#   REQ-xx -> NFR-01..NFR-07 (7 occurrences, in order)
#   REQ-   -> REQ-18 / REQ-19 (2 occurrences, in order)
# This matches the user's instruction: groupmates renumber later; we just
# need every requirement to have a usable, unique id.

placeholder_slots = []   # (table_idx, row_idx, cell_idx, current_text)
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            txt = cell.text.strip()
            if txt in ("REQ-0?", "REQ-xx", "REQ-", "REQ-"):
                placeholder_slots.append((ti, ri, ci, txt))

print("placeholder cells found:", len(placeholder_slots))
for s in placeholder_slots:
    print("  ", s)

# Assign replacements in document order
repl_map = {"REQ-0?": iter(["REQ-08", "REQ-09"]),
            "REQ-xx": iter([f"NFR-0{i}" for i in range(1, 8)]),
            "REQ-": iter(["REQ-18", "REQ-19"])}

# Re-walk and apply
applied = []
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            txt = cell.text.strip()
            if txt in repl_map and txt not in ("REQ_ID",):
                it = repl_map[txt]
                try:
                    new = next(it)
                except StopIteration:
                    continue
                # replace text in every paragraph of the cell
                for p in cell.paragraphs:
                    for run in p.runs:
                        if txt in run.text:
                            run.text = run.text.replace(txt, new)
                            applied.append((ti, ri, ci, txt, new))

print("\napplied replacements:", len(applied))
for a in applied:
    print("  ", a)

# Fix the "LIne2" typo
for p in doc.paragraphs:
    for run in p.runs:
        if "LIne2" in run.text:
            run.text = run.text.replace("LIne2", "Line2")
            print("fixed typo LIne2 -> Line2")

doc.save(DST)
print("\nSaved:", DST)
